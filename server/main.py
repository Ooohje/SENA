"""
SENA Backend — FastAPI server
RTX 5070 Ti · WavLM scoring · Faster-Whisper STT · Ollama Qwen3-8B · Kokoro TTS
"""
import os, json, re, uuid, base64, io, subprocess
from pathlib import Path
from contextlib import asynccontextmanager
from typing import Optional
from datetime import datetime

import torch
import numpy as np
import librosa
import torch.nn as nn
from transformers import AutoModel, AutoFeatureExtractor
from faster_whisper import WhisperModel
import httpx
import soundfile as sf

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
import uvicorn

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Kokoro TTS (optional — graceful fallback if not installed) ────────────────
try:
    from kokoro import KPipeline as _KokoroPipeline
    _KOKORO_OK = True
except ImportError:
    _KOKORO_OK = False
    print("[WARN] kokoro not installed → TTS disabled (pip install kokoro)")

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR   = Path(__file__).parent
PROJ_ROOT  = BASE_DIR.parent
WEIGHTS    = PROJ_ROOT / "best_model" / "wavlm_738_755.pt"
TMP_DIR    = BASE_DIR / "audio_tmp"
RPT_DIR    = BASE_DIR / "report_tmp"
TMP_DIR.mkdir(exist_ok=True)
RPT_DIR.mkdir(exist_ok=True)

# ── Constants ─────────────────────────────────────────────────────────────────
BACKBONE     = "microsoft/wavlm-base"
OLLAMA_BASE  = "http://localhost:11434"
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen3:8b")

# ── Global model registry ─────────────────────────────────────────────────────
G: dict = {}       # evaluator, stt, tts
SESSIONS: dict = {}  # sessionId → {report, history, lang}


# ═══════════════════════════════════════════════════════════════════════════════
# SCORING MODEL
# ═══════════════════════════════════════════════════════════════════════════════

class MultiTaskAudioRegression(nn.Module):
    def __init__(self, model_name: str = BACKBONE, dropout_prob: float = 0.1):
        super().__init__()
        self.backbone = AutoModel.from_pretrained(model_name)
        h = self.backbone.config.hidden_size
        def _head():
            return nn.Sequential(
                nn.Dropout(dropout_prob),
                nn.Linear(h, 256), nn.ReLU(), nn.Linear(256, 1)
            )
        self.artic_head  = _head()
        self.prosody_head = _head()

    def forward(self, input_values, attention_mask=None):
        out    = self.backbone(input_values=input_values, attention_mask=attention_mask)
        pooled = torch.mean(out.last_hidden_state, dim=1)
        return self.artic_head(pooled), self.prosody_head(pooled)


class SpeechEvaluator:
    SR      = 16_000
    MAX_LEN = 16_000 * 8   # 8 s
    PAD     = 16_000 // 4  # 250 ms

    def __init__(self):
        # RTX 5070 Ti (Blackwell sm_120) not yet supported by PyTorch CUDA kernels → CPU
        self.dev = torch.device("cpu")
        print(f"[Evaluator] backbone={BACKBONE}  device={self.dev}")

        self.processor = AutoFeatureExtractor.from_pretrained(BACKBONE)

        print("[Evaluator] loading Silero VAD…")
        self.vad, utils = torch.hub.load(
            "snakers4/silero-vad", "silero_vad", force_reload=False
        )
        self.get_ts = utils[0]
        self.vad.to("cpu")  # Silero VAD stays on CPU (RTX 5070 Ti sm_120 not supported)

        print(f"[Evaluator] loading weights from {WEIGHTS}…")
        self.model = MultiTaskAudioRegression().to(self.dev)
        self.model.load_state_dict(torch.load(str(WEIGHTS), map_location=self.dev))
        self.model.eval()
        print("[Evaluator] ✅ ready")

    def _preprocess(self, path: str) -> np.ndarray:
        wav, _ = librosa.load(path, sr=self.SR, mono=True, dtype=np.float32)
        wav    = librosa.util.normalize(wav).astype(np.float32)
        t      = torch.from_numpy(wav)  # VAD always runs on CPU
        ts     = self.get_ts(t, self.vad, sampling_rate=self.SR, threshold=0.5)
        if ts:
            s   = max(0, ts[0]["start"] - self.PAD)
            e   = min(len(wav), ts[-1]["end"] + self.PAD)
            wav = wav[s:e]
        return wav

    def predict(self, path: str) -> dict:
        wav  = self._preprocess(path)
        inp  = self.processor(
            wav, sampling_rate=self.SR, return_tensors="pt",
            padding="max_length", max_length=self.MAX_LEN, truncation=True
        )
        with torch.no_grad():
            a, p = self.model(inp.input_values.to(self.dev))
        art = round(max(1.0, min(5.0, float(a.item()))), 2)
        pro = round(max(1.0, min(5.0, float(p.item()))), 2)
        return {"articulation": art, "prosody": pro, "overall": round((art + pro) / 2, 2)}


# ═══════════════════════════════════════════════════════════════════════════════
# STT — Faster-Whisper
# ═══════════════════════════════════════════════════════════════════════════════

class STTEngine:
    def __init__(self, size: str = "medium"):
        dev   = "cuda" if torch.cuda.is_available() else "cpu"
        ctype = "float16" if dev == "cuda" else "int8"
        print(f"[STT] loading faster-whisper {size} on {dev}…")
        self.model = WhisperModel(size, device=dev, compute_type=ctype)
        print("[STT] ✅ ready")

    def transcribe(self, path: str) -> str:
        segs, _ = self.model.transcribe(path, language="en", beam_size=5, vad_filter=True)
        return " ".join(s.text.strip() for s in segs).strip()


# ═══════════════════════════════════════════════════════════════════════════════
# TTS — Kokoro
# ═══════════════════════════════════════════════════════════════════════════════

class TTSEngine:
    SR    = 24_000
    VOICE = "af_heart"

    def __init__(self):
        print("[TTS] loading Kokoro…")
        try:
            self.pipe = _KokoroPipeline(lang_code="a")
        except RuntimeError as e:
            if "CUDA" in str(e):
                print("[TTS] CUDA not supported for Kokoro on this GPU → falling back to CPU")
                self.pipe = _KokoroPipeline(lang_code="a", device="cpu")
            else:
                raise
        print("[TTS] ✅ ready")

    def synth_b64(self, text: str) -> str:
        """Return data:audio/wav;base64,... string."""
        chunks = [audio for _, _, audio in
                  self.pipe(text, voice=self.VOICE, speed=1.0, split_pattern=r"\n+")]
        if not chunks:
            return ""
        arr = np.concatenate(chunks)
        buf = io.BytesIO()
        sf.write(buf, arr, self.SR, format="WAV", subtype="PCM_16")
        return "data:audio/wav;base64," + base64.b64encode(buf.getvalue()).decode()


# ═══════════════════════════════════════════════════════════════════════════════
# LLM — Ollama Qwen3-8B
# ═══════════════════════════════════════════════════════════════════════════════

def _strip_think(text: str) -> str:
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()

async def _ollama(messages: list, think: bool = False, max_tokens: int = 400) -> str:
    msgs = [dict(m) for m in messages]
    if not think:
        # Disable Qwen3 thinking mode for real-time turns
        for m in msgs:
            if m["role"] == "system":
                m["content"] = m["content"].rstrip() + "\n/no_think"
                break
    async with httpx.AsyncClient(timeout=180.0) as c:
        r = await c.post(
            f"{OLLAMA_BASE}/api/chat",
            json={
                "model": OLLAMA_MODEL,
                "messages": msgs,
                "stream": False,
                "options": {"temperature": 0.7 if not think else 0.3, "num_predict": max_tokens},
            },
        )
        r.raise_for_status()
        return _strip_think(r.json()["message"]["content"])

_LEVELS = {"beg": "beginner", "int": "intermediate", "adv": "advanced"}
_TOPIC_STYLE = {
    "opic":      "Ask the learner to describe personal experiences, routines, and opinions — the kind of open-ended questions used in spoken English assessments.",
    "free":      "Have a relaxed, casual chat about whatever is on the learner's mind.",
    "biz":       "Simulate professional workplace situations: status updates, meetings, emails, negotiations.",
    "travel":    "Roleplay travel situations: airports, hotels, asking for directions, ordering food.",
    "interview": "Conduct a natural English job interview: background, strengths, experience, goals.",
    "daily":     "Chat about everyday topics: food, hobbies, weekend plans, daily routines.",
}

def _chat_system(level: str, topic: str) -> str:
    style = _TOPIC_STYLE.get(topic, _TOPIC_STYLE["free"])
    return (
        f"You are SENA, a friendly AI English tutor for Korean learners. "
        f"The learner's level is {_LEVELS.get(level, 'intermediate')}. "
        f"{style} "
        "Respond in natural, clear English. Keep each reply to 2-3 sentences "
        "and always end with one short, engaging follow-up question. "
        "Do NOT announce the topic name or exam name. "
        "Do not use markdown or bullet points."
    )


# ═══════════════════════════════════════════════════════════════════════════════
# WORD REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def _build_word(report: dict, history: list, lang: str) -> Path:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    h = doc.add_heading(
        "SENA AI 튜터 · 세션 리포트" if lang == "ko" else "SENA AI Tutor · Session Report", 0
    )
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"{'생성일' if lang=='ko' else 'Generated'}: {datetime.now().strftime('%Y-%m-%d  %H:%M')}"
    )
    doc.add_paragraph()

    # 1. Overall
    doc.add_heading("📊 총평" if lang == "ko" else "📊 Overall Assessment", 1)
    doc.add_paragraph(report.get("overall", ""))

    # 2. Grammar
    grammar = report.get("grammar") or []
    if grammar:
        doc.add_heading("📝 문법 교정" if lang == "ko" else "📝 Grammar Corrections", 1)
        for err in grammar:
            p = doc.add_paragraph(); p.add_run(f"• {err.get('issue','')}")
            p = doc.add_paragraph()
            p.add_run(("  예시: " if lang=="ko" else "  Example: ")).bold = True
            p.add_run(err.get("example",""))
            p = doc.add_paragraph()
            r2 = p.add_run("  ✅ 수정: " if lang=="ko" else "  ✅ Correction: "); r2.bold = True
            p.add_run(err.get("correction",""))
            doc.add_paragraph()

    # 3. Vocabulary
    vocab = report.get("vocabulary") or []
    if vocab:
        doc.add_heading("💬 어휘 제안" if lang == "ko" else "💬 Vocabulary Suggestions", 1)
        for v in vocab:
            doc.add_paragraph(f"→ {v.get('suggestion','')}", style="List Bullet")

    # 4. Fluency
    fluency = report.get("fluency") or {}
    if fluency.get("notes"):
        doc.add_heading("🎯 유창성" if lang == "ko" else "🎯 Fluency", 1)
        doc.add_paragraph(fluency["notes"])

    # 5. Transcript
    doc.add_heading("📋 대화 기록" if lang == "ko" else "📋 Session Transcript", 1)
    for m in history:
        is_ai  = m.get("who") == "ai"
        label  = ("AI 튜터" if lang=="ko" else "AI Tutor") if is_ai else ("학습자" if lang=="ko" else "Learner")
        p      = doc.add_paragraph()
        p.add_run(f"[{label}]  ").bold = True
        p.add_run(m.get("text",""))
        sc = m.get("score")
        if sc:
            sp = doc.add_paragraph()
            sp.add_run(f"     발음: Art {sc['articulation']:.1f}  Pro {sc['prosody']:.1f}").italic = True

    path = RPT_DIR / f"report_{uuid.uuid4().hex[:8]}.docx"
    doc.save(str(path))
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _to_wav(src: Path, dst: Path):
    r = subprocess.run(
        ["ffmpeg", "-y", "-i", str(src), "-ar", "16000", "-ac", "1", "-f", "wav", str(dst)],
        capture_output=True,
    )
    if r.returncode != 0:
        raise RuntimeError(f"ffmpeg failed: {r.stderr.decode()}")

async def _save_wav(file: UploadFile) -> Path:
    data = await file.read()
    if not data:
        raise HTTPException(400, "Empty audio file")
    stem = uuid.uuid4().hex[:10]
    raw  = TMP_DIR / f"{stem}_in"
    wav  = TMP_DIR / f"{stem}.wav"
    raw.write_bytes(data)
    try:
        _to_wav(raw, wav)
    finally:
        raw.unlink(missing_ok=True)
    return wav

def _feedback(art: float, pro: float, lang: str) -> list[str]:
    def pick(v, hi, mid, lo):
        return hi if v >= 4.3 else mid if v >= 3.8 else lo
    if lang == "ko":
        return [
            pick(art,
                 "발음이 매우 정확합니다.",
                 "발음은 양호하나 /s/, /th/, /r/ 발음을 더 연습해보세요.",
                 "자음·모음 발음에 집중 연습이 필요합니다."),
            pick(pro,
                 "운율과 억양이 매우 자연스럽습니다.",
                 "운율은 자연스러우나 내용어 강세를 더 강조해보세요.",
                 "강세·리듬·인토네이션을 원어민 음성을 들으며 연습해보세요."),
            pick((art+pro)/2,
                 "훌륭한 발화입니다! 이 수준을 유지해보세요.",
                 "꾸준히 연습하면 더욱 향상될 것입니다.",
                 "반복 연습을 통해 충분히 향상될 수 있습니다."),
        ]
    return [
        pick(art, "Excellent articulation!", "Good — refine /s/, /th/, /r/.", "Focus on pronouncing each phoneme clearly."),
        pick(pro, "Great prosody and intonation!", "Natural rhythm — stress content words more.", "Practice stress and intonation by listening to native speakers."),
        pick((art+pro)/2, "Excellent speaking — keep it up!", "Consistent practice will take you further.", "With regular practice you'll see clear improvement."),
    ]


# ═══════════════════════════════════════════════════════════════════════════════
# APP
# ═══════════════════════════════════════════════════════════════════════════════

@asynccontextmanager
async def lifespan(app: FastAPI):
    print("\n" + "═"*48)
    print("  SENA Backend  —  model loading…")
    print("═"*48)
    G["evaluator"] = SpeechEvaluator()
    G["stt"]       = STTEngine()
    G["tts"]       = TTSEngine() if _KOKORO_OK else None
    print("═"*48)
    print("  ✅  All models ready.  Server listening on :8000")
    print("═"*48 + "\n")
    yield
    # cleanup old reports on shutdown
    for f in RPT_DIR.glob("*.docx"):
        f.unlink(missing_ok=True)

app = FastAPI(title="SENA API", version="1.0.0", lifespan=lifespan)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

FRONTEND_DIR = PROJ_ROOT  # HTML/JS/CSS files are at the project root


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {
        "status": "ok",
        "gpu": torch.cuda.is_available(),
        "gpu_name": torch.cuda.get_device_name(0) if torch.cuda.is_available() else None,
        "tts": G.get("tts") is not None,
        "models": list(G.keys()),
    }


@app.post("/api/v1/score")
async def score_pronunciation(
    audio:     UploadFile = File(...),
    reference: str        = Form(""),
    lang:      str        = Form("ko"),
):
    wav = await _save_wav(audio)
    user_text = ""
    try:
        result = G["evaluator"].predict(str(wav))
        try:
            user_text = G["stt"].transcribe(str(wav))
        except Exception:
            user_text = ""
    finally:
        wav.unlink(missing_ok=True)
    result["feedback"] = _feedback(result["articulation"], result["prosody"], lang)
    result["userText"] = user_text
    return result


@app.post("/api/v1/chat/turn")
async def chat_turn(
    level:   str                    = Form("int"),
    topic:   str                    = Form("opic"),
    history: str                    = Form("[]"),
    lang:    str                    = Form("ko"),
    audio:   Optional[UploadFile]   = File(None),
):
    hist        = json.loads(history)
    user_text   = ""
    utter_score = None

    # STT + scoring
    if audio and audio.filename:
        data = await audio.read()
        if data:
            stem = uuid.uuid4().hex[:10]
            raw  = TMP_DIR / f"{stem}_in"
            wav  = TMP_DIR / f"{stem}.wav"
            raw.write_bytes(data)
            try:
                _to_wav(raw, wav)
                user_text   = G["stt"].transcribe(str(wav))
                sc          = G["evaluator"].predict(str(wav))
                utter_score = {"articulation": sc["articulation"], "prosody": sc["prosody"]}
            except Exception as e:
                print(f"[chat/turn] STT/score error: {e}")
            finally:
                raw.unlink(missing_ok=True)
                wav.unlink(missing_ok=True)

    # Build messages for Ollama — accept both {who,text} and {role,content} formats
    messages = [{"role": "system", "content": _chat_system(level, topic)}]
    for m in hist:
        if "role" in m:
            messages.append({"role": m["role"], "content": m.get("content", m.get("text", ""))})
        else:
            messages.append({"role": "assistant" if m.get("who") == "ai" else "user", "content": m.get("text", "")})
    if user_text:
        messages.append({"role": "user", "content": user_text})

    ai_text = await _ollama(messages, think=False, max_tokens=300)

    # TTS
    ai_audio_b64 = ""
    if G.get("tts") and ai_text:
        try:
            ai_audio_b64 = G["tts"].synth_b64(ai_text)
        except Exception as e:
            print(f"[TTS] error: {e}")

    return {
        "userText":       user_text,
        "aiText":         ai_text,
        "aiAudioBase64":  ai_audio_b64,
        "utteranceScore": utter_score,
    }


@app.post("/api/v1/chat/reset")
async def chat_reset():
    return {"ok": True}


@app.post("/api/v1/report/generate")
async def report_generate(payload: dict):
    history = payload.get("history", [])
    lang    = payload.get("lang", "ko")

    if not history:
        raise HTTPException(400, "history is empty")

    # Build transcript
    lines = []
    for m in history:
        is_ai = m.get("who") == "ai"
        label = "AI Tutor" if is_ai else "Student"
        sc    = m.get("score")
        sc_s  = f" [Art:{sc['articulation']:.1f} Pro:{sc['prosody']:.1f}]" if sc else ""
        lines.append(f"{label}{sc_s}: {m.get('text','')}")
    transcript = "\n".join(lines)

    if lang == "ko":
        sys_prompt = (
            "당신은 영어 교육 전문가입니다. 아래 학습자의 대화를 분석하고 "
            "다음 JSON 형식으로만 응답하세요 (마크다운 없이):\n"
            '{\n'
            '  "overall": "전반적 평가 2-3문장",\n'
            '  "grammar": [{"issue":"오류설명","example":"오류문장","correction":"수정문장"}],\n'
            '  "vocabulary": [{"suggestion":"더 자연스러운 표현 제안"}],\n'
            '  "fluency": {"notes":"유창성 피드백 2-3문장"}\n'
            '}'
        )
        user_msg = f"대화 기록:\n\n{transcript}"
    else:
        sys_prompt = (
            "You are an expert English teacher. Analyze the learner's conversation and "
            "return ONLY valid JSON (no markdown fences):\n"
            '{\n'
            '  "overall": "overall assessment 2-3 sentences",\n'
            '  "grammar": [{"issue":"error description","example":"erroneous sentence","correction":"corrected version"}],\n'
            '  "vocabulary": [{"suggestion":"more natural expression suggestion"}],\n'
            '  "fluency": {"notes":"fluency feedback 2-3 sentences"}\n'
            '}'
        )
        user_msg = f"Conversation:\n\n{transcript}"

    raw = await _ollama(
        [{"role": "system", "content": sys_prompt}, {"role": "user", "content": user_msg}],
        think=True, max_tokens=1400,
    )

    report: dict = {}
    try:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            report = json.loads(m.group())
    except Exception:
        report = {"overall": raw, "grammar": [], "vocabulary": [], "fluency": {"notes": ""}}

    sid = uuid.uuid4().hex[:12]
    SESSIONS[sid] = {"report": report, "history": history, "lang": lang}
    return {"sessionId": sid, "report": report}


@app.get("/api/v1/report/download/{session_id}")
async def report_download(session_id: str):
    sess = SESSIONS.get(session_id)
    if not sess:
        raise HTTPException(404, "Session not found — it may have expired. Please regenerate the report.")
    path  = _build_word(sess["report"], sess["history"], sess["lang"])
    fname = "SENA_세션리포트.docx" if sess["lang"] == "ko" else "SENA_Session_Report.docx"
    return FileResponse(
        str(path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=fname,
    )


# ── Frontend (must be mounted AFTER all API routes) ───────────────────────────
@app.get("/")
async def _root():
    return FileResponse(FRONTEND_DIR / "index.html")

@app.get("/{page}.html")
async def _page(page: str):
    f = FRONTEND_DIR / f"{page}.html"
    if not f.exists():
        raise HTTPException(status_code=404)
    return FileResponse(f)

app.mount("/", StaticFiles(directory=str(FRONTEND_DIR)), name="static")


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=False)
